"""
Document processing utilities for the Bank AI Chatbot.
This module handles the extraction of text from various document formats,
chunking the text into manageable pieces, and creating embeddings.
"""

import os
import re
import json
import subprocess
from pathlib import Path
from typing import List, Dict, Any, Tuple

import nltk
from nltk.tokenize import sent_tokenize
from sentence_transformers import SentenceTransformer

# Download necessary NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

from src.config import (
    DOCUMENTS_DIR, 
    PROCESSED_DIR, 
    EMBEDDINGS_DIR, 
    CHUNK_SIZE, 
    CHUNK_OVERLAP,
    EMBEDDING_MODEL
)

class DocumentProcessor:
    """Process documents and create embeddings for the knowledge base."""
    
    def __init__(self):
        """Initialize the document processor."""
        self.embedding_model = SentenceTransformer(EMBEDDING_MODEL)
    
    def process_all_documents(self) -> Dict[str, int]:
        """
        Process all documents in the documents directory.
        
        Returns:
            Dict[str, int]: A dictionary with document categories as keys and 
                           the number of processed documents as values.
        """
        stats = {}
        
        # Process each category directory
        for category_dir in DOCUMENTS_DIR.iterdir():
            if category_dir.is_dir():
                category_name = category_dir.name
                processed_count = self._process_category(category_dir, category_name)
                stats[category_name] = processed_count
        
        return stats
    
    def _process_category(self, category_dir: Path, category_name: str) -> int:
        """
        Process all documents in a category directory.
        
        Args:
            category_dir (Path): Path to the category directory.
            category_name (str): Name of the category.
            
        Returns:
            int: Number of processed documents.
        """
        processed_count = 0
        
        # Create output directories if they don't exist
        processed_category_dir = PROCESSED_DIR / category_name
        embeddings_category_dir = EMBEDDINGS_DIR / category_name
        
        os.makedirs(processed_category_dir, exist_ok=True)
        os.makedirs(embeddings_category_dir, exist_ok=True)
        
        # Process each document in the category
        for doc_path in category_dir.iterdir():
            if doc_path.is_file():
                try:
                    # Extract text from document
                    text = self._extract_text(doc_path)
                    
                    # Skip empty documents
                    if not text.strip():
                        print(f"Warning: Empty document {doc_path}")
                        continue
                    
                    # Create chunks
                    chunks = self._create_chunks(text)
                    
                    # Save processed text
                    processed_path = processed_category_dir / f"{doc_path.stem}.json"
                    with open(processed_path, 'w') as f:
                        json.dump({
                            'source': str(doc_path),
                            'chunks': chunks
                        }, f, indent=2)
                    
                    # Create and save embeddings
                    self._create_embeddings(chunks, doc_path.stem, embeddings_category_dir)
                    
                    processed_count += 1
                    
                except Exception as e:
                    print(f"Error processing {doc_path}: {e}")
        
        return processed_count
    
    def _extract_text(self, doc_path: Path) -> str:
        """
        Extract text from a document.
        
        Args:
            doc_path (Path): Path to the document.
            
        Returns:
            str: Extracted text.
        """
        if doc_path.suffix.lower() == '.pdf':
            return self._extract_text_from_pdf(doc_path)
        elif doc_path.suffix.lower() == '.docx':
            return self._extract_text_from_docx(doc_path)
        elif doc_path.suffix.lower() in ['.txt', '.md']:
            return self._extract_text_from_text_file(doc_path)
        else:
            raise ValueError(f"Unsupported file format: {doc_path.suffix}")
    
    def _extract_text_from_pdf(self, pdf_path: Path) -> str:
        """
        Extract text from a PDF file using pdftotext.
        
        Args:
            pdf_path (Path): Path to the PDF file.
            
        Returns:
            str: Extracted text.
        """
        try:
            # Use pdftotext from poppler-utils
            result = subprocess.run(
                ['pdftotext', str(pdf_path), '-'],
                capture_output=True,
                text=True,
                check=True
            )
            return result.stdout
        except subprocess.CalledProcessError:
            # Fallback to pdf2image and OCR if needed
            print(f"Warning: Could not extract text from {pdf_path} using pdftotext")
            return f"[PDF EXTRACTION FAILED: {pdf_path}]"
    
    def _extract_text_from_docx(self, docx_path: Path) -> str:
        
        """
        Extract text from a DOCX file.
        
        Args:
            docx_path (Path): Path to the DOCX file.
            
        Returns:
            str: Extracted text.
        """
        try:
            import docx # type: ignore
            doc = docx.Document(docx_path)
            return '\n\n'.join([para.text for para in doc.paragraphs])
        except ImportError:
            print("Warning: python-docx not installed. Installing...")
            subprocess.run(['pip', 'install', 'python-docx'], check=True)
            import docx
            doc = docx.Document(docx_path)
            return '\n\n'.join([para.text for para in doc.paragraphs])
    
    def _extract_text_from_text_file(self, text_path: Path) -> str:
        
        
        
        """
        Extract text from a text file.
        
        Args:
            text_path (Path): Path to the text file.
            
        Returns:
            str: Extracted text.
        """
        with open(text_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """
        Split text into chunks with metadata.
        
        Args:
            text (str): Text to split.
            
        Returns:
            List[Dict[str, Any]]: List of chunks with metadata.
        """
        # Clean text
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Split into sentences
        sentences = sent_tokenize(text)
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            # If adding this sentence would exceed chunk size, save current chunk and start a new one
            if len(current_chunk) + len(sentence) > CHUNK_SIZE and current_chunk:
                chunks.append({
                    'text': current_chunk.strip(),
                    'size': len(current_chunk)
                })
                
                # Start new chunk with overlap
                words = current_chunk.split()
                overlap_words = words[-int(CHUNK_OVERLAP/10):] if len(words) > int(CHUNK_OVERLAP/10) else []
                current_chunk = ' '.join(overlap_words) + ' ' + sentence
            else:
                # Add sentence to current chunk
                current_chunk += ' ' + sentence
        
        # Add the last chunk if it's not empty
        if current_chunk.strip():
            chunks.append({
                'text': current_chunk.strip(),
                'size': len(current_chunk)
            })
        
        return chunks
    
    def _create_embeddings(self, chunks: List[Dict[str, Any]], doc_name: str, output_dir: Path) -> None:
        """
        Create embeddings for chunks and save them.
        
        Args:
            chunks (List[Dict[str, Any]]): List of text chunks.
            doc_name (str): Document name.
            output_dir (Path): Output directory.
        """
        # Extract text from chunks
        texts = [chunk['text'] for chunk in chunks]
        
        # Create embeddings
        embeddings = self.embedding_model.encode(texts)
        
        # Save embeddings
        embeddings_path = output_dir / f"{doc_name}.json"
        with open(embeddings_path, 'w') as f:
            json.dump({
                'document': doc_name,
                'embeddings': embeddings.tolist()
            }, f)


if __name__ == "__main__":
    processor = DocumentProcessor()
    stats = processor.process_all_documents()
    print(f"Processed documents: {stats}")
